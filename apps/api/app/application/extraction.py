# SPDX-FileCopyrightText: 2026 Keresztes Zsolt <https://kereszteszsolt.hu>
# SPDX-License-Identifier: Apache-2.0

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from ..rag.native.chunking import TextSection

SUPPORTED_EXTRACTION_SUFFIXES = frozenset({".pdf", ".docx", ".txt", ".md", ".markdown"})


def extract_sections(path: Path) -> list[TextSection]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix in {".txt", ".md", ".markdown"}:
        return [TextSection(path.read_text(encoding="utf-8", errors="replace"))]
    raise ValueError(f"Unsupported document type: {suffix}")


def _extract_pdf(path: Path) -> list[TextSection]:
    reader = PdfReader(str(path))
    if reader.is_encrypted:
        raise ValueError("Encrypted PDFs are not supported.")
    return [
        TextSection(text=page.extract_text() or "", page_number=index + 1)
        for index, page in enumerate(reader.pages)
    ]


def _extract_docx(path: Path) -> list[TextSection]:
    document = DocxDocument(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        if cell.text.strip()
    ]
    return [TextSection(text="\n\n".join([*paragraphs, *table_cells]))]
