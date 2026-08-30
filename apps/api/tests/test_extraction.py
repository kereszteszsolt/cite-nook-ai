# SPDX-FileCopyrightText: 2026 Keresztes Zsolt <https://kereszteszsolt.hu>
# SPDX-License-Identifier: Apache-2.0

from pathlib import Path

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from app.application.extraction import extract_sections


@pytest.mark.parametrize("suffix", [".txt", ".md", ".markdown"])
def test_plain_text_and_markdown_extraction(tmp_path: Path, suffix: str) -> None:
    path = tmp_path / f"notes{suffix}"
    path.write_text("# Heading\n\nUseful content", encoding="utf-8")

    sections = extract_sections(path)

    assert [section.text for section in sections] == ["# Heading\n\nUseful content"]
    assert sections[0].page_number is None


def test_docx_extracts_paragraphs_and_table_cells(tmp_path: Path) -> None:
    path = tmp_path / "report.docx"
    document = DocxDocument()
    document.add_paragraph("Paragraph content")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Table content"
    document.save(path)

    sections = extract_sections(path)

    assert len(sections) == 1
    assert "Paragraph content" in sections[0].text
    assert "Table content" in sections[0].text


def test_pdf_extraction_retains_one_based_page_numbers(
    tmp_path: Path,
) -> None:
    path = tmp_path / "report.pdf"
    writer = PdfWriter()
    _add_pdf_text_page(writer, "First page")
    _add_pdf_text_page(writer, "Second page")
    with path.open("wb") as output:
        writer.write(output)

    sections = extract_sections(path)

    assert [(section.text, section.page_number) for section in sections] == [
        ("First page", 1),
        ("Second page", 2),
    ]


def _add_pdf_text_page(writer: PdfWriter, text: str) -> None:
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font})}
    )
    content = DecodedStreamObject()
    content.set_data(f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode())
    page[NameObject("/Contents")] = content
